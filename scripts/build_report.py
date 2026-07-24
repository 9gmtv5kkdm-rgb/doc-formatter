# -*- coding: utf-8 -*-
"""普通报告 / 总结排版参考脚本（非红头公文）。
本技能原 build_doc.py 只覆盖 GB/T 9704 红头公文。本脚本补上触发词承诺的
"自动生成目录 + 标题层级"能力——把公文层级文字映射为 Word 标题样式（Heading 1-4），
并插入可更新的目录域（TOC）。使用前替换 SRC / OUT / 内容文本。

依赖：python-docx。字体选择沿用 doc_style.py 的本机实测规则。
"""
import re
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from doc_style import *   # FONT_FS/FONT_HEI/FONT_KAI/FONT_SONG/PT3..PT4

SRC = r"【源文档路径.docx】"   # 已有内容，可选
OUT = r"【目标输出路径.docx】"

# ===================== 1. 标题层级 → Word 标题样式映射 =====================
# 公文结构序数：一、→（一）→1.→（1）。映射到 Heading 1~4，使目录自动抓取。
LEVEL_PATTERNS = [
    (re.compile(r"^[一二三四五六七八九十]+、"), "Heading 1"),   # 一、
    (re.compile(r"^（[一二三四五六七八九十]+）"), "Heading 2"),  # （一）
    (re.compile(r"^\d+\."), "Heading 3"),                       # 1.
    (re.compile(r"^（\d+）"), "Heading 4"),                     # （1）
]

def assign_heading_style(para):
    """根据段落文本首部序数判断层级并套用 Word 标题样式；返回层级或 None。"""
    t = para.text.strip()
    for pat, style_name in LEVEL_PATTERNS:
        if pat.match(t):
            para.style = para.document.styles[style_name]
            # 按公文规范定字体：一级黑体、二级楷体、三级仿宋
            if style_name == "Heading 1":
                set_east_asia(para, FONT_HEI)
            elif style_name == "Heading 2":
                set_east_asia(para, FONT_KAI)
            else:
                set_east_asia(para, FONT_FS)
            return style_name
    return None

def set_east_asia(para, ea):
    for r in para.runs:
        rPr = r._element.get_or_add_rPr()
        rf = rPr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts"); rPr.append(rf)
        rf.set(qn("w:ascii"), ea); rf.set(qn("w:hAnsi"), ea); rf.set(qn("w:eastAsia"), ea)

# ===================== 2. 插入自动目录域（TOC） =====================
def add_toc_field(document, max_level=3):
    """在文档开头插入 TOC 域。注意：域需在 Word 中右键「更新域」才会填充，
    python-docx 无法在生成时计算页码，故预留占位提示。"""
    placeholder = document.add_paragraph()
    run = placeholder.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = f'TOC \\o "1-{max_level}" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    tip = OxmlElement("w:t"); tip.text = "（在 Word 中右键此区 → 更新域，生成目录）"
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin); run._element.append(instr)
    run._element.append(fld_sep); run._element.append(tip); run._element.append(fld_end)
    return placeholder

# ===================== 3. 页眉 / 页码（普通报告） =====================
def add_page_number_footer(section, align=WD_ALIGN_PARAGRAPH.CENTER):
    """普通报告页码：居中、宋体四号半角阿拉伯数字（无奇偶页差异要求时用此简化版）。"""
    p = section.footer.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(14)
    r = p.add_run("— "); set_east_asia_run(r, FONT_SONG, PT4)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    r2 = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts"); rf.set(qn("w:ascii"), FONT_SONG); rf.set(qn("w:hAnsi"), FONT_SONG); rf.set(qn("w:eastAsia"), FONT_SONG)
    rPr.append(rf); sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(PT4*2)); rPr.append(sz)
    r2.append(rPr); r2.append(OxmlElement("w:t")); fld.append(r2)
    p._p.append(fld)
    r3 = p.add_run(" —"); set_east_asia_run(r3, FONT_SONG, PT4)

def set_east_asia_run(r, ea, size):
    r.font.size = Pt(size)
    rPr = r._element.get_or_add_rPr()
    rf = OxmlElement("w:rFonts"); rf.set(qn("w:ascii"), ea); rf.set(qn("w:hAnsi"), ea); rf.set(qn("w:eastAsia"), ea)
    rPr.append(rf)

# ===================== 4. 用法示例（按需改写） =====================
if __name__ == "__main__":
    doc = Document()
    # 页面 A4（普通报告常用边距：上下 25mm / 左右 25mm，可依单位模板调整）
    sec = doc.sections[0]
    sec.page_width = Mm(210); sec.page_height = Mm(297)
    sec.top_margin = Mm(25); sec.bottom_margin = Mm(25)
    sec.left_margin = Mm(25); sec.right_margin = Mm(25)
    # 默认正文字体仿宋三号
    normal = doc.styles["Normal"]; normal.font.name = FONT_FS; normal.font.size = Pt(PT3)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FS)
    # 1) 插入目录域
    add_toc_field(doc, max_level=3)
    # 2) 写标题层级（自动套 Heading 样式）
    for title, body in [
        ("一、工作总体情况", "本年度围绕核心目标推进各项工作，取得阶段性成效。"),
        ("（一）主要指标", "营收同比增长 12%，客户满意度提升至 94%。"),
        ("1. 重点举措", "优化流程、加强协同，缩短交付周期。"),
    ]:
        h = doc.add_paragraph(title); assign_heading_style(h)
        b = doc.add_paragraph(body); b.paragraph_format.first_line_indent = Pt(PT3*2)
    # 3) 页码
    add_page_number_footer(sec)
    doc.save(OUT)
    print("SAVED:", OUT)
