# -*- coding: utf-8 -*-
"""按 GB/T 9704-2012 公文格式重建 docx。
使用前请修改：
1. OUT → 目标输出路径
2. SRC → 源文档路径（若作为样式参考）
3. 第 1 步-第 3 步中的 add_para() 文本 → 替换为实际内容
"""
import os
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"【目标输出路径.docx】"
SRC = r"【源文档路径.docx】"
# 示例：
# OUT = r"C:\Users\用户名\Desktop\XXXX公司（公文重排版）.docx"
# SRC = r"C:\Users\用户名\Desktop\XXXX公司.docx"

# ===================== 公文身份配置（替换为你的机关 / 文件） =====================
# 定位说明：本脚本是「标准公文范本 / 演示生成器」（示例：国有资产清查盘活方案），
# 用于配合 verify_doc.py 演示 GB/T 9704 格式校验。改下方常量仅替换红头/标题/署名，
# 正文（一~三级标题、人员名单等）为示例样本，生成其他机关公文需自行改写。
ORG_NAME = "河南九天封头制造有限公司"            # 发文机关标志（红头，示例值）
DOC_TITLE = "国有资产清查盘活专项工作实施方案"    # 公文标题，示例值
LEADER_GROUP = "河南九天封头制造有限公司国有资产清查盘活工作领导小组人员名单"  # 附件标题，示例值
DOC_DATE = "2026年7月7日"                         # 成文日期，示例值，按需替换

# ---------- 字体/字号常量（来自共享模块，确保与 verify 一致） ----------
from doc_style import *  # FONT_XBS/FONT_FS/FONT_HEI/FONT_KAI/FONT_SONG/RED/PT3..PT4/FIRSTLINE_COEF

doc = Document()

# 默认正文字体
style = doc.styles["Normal"]
style.font.name = FONT_FS
style.font.size = Pt(PT3)
style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FS)

# ---------- 页面与版心 (P1-P5) ----------
sec = doc.sections[0]
sec.page_width = Mm(210)
sec.page_height = Mm(297)
sec.top_margin = Mm(37)
sec.bottom_margin = Mm(35)
sec.left_margin = Mm(28)
sec.right_margin = Mm(26)
# 页脚距页底 28mm（一字线上距版心下边缘 7mm；版心下边缘距底 35mm）PG4
secG = sec._sectPr.find(qn("w:pgMar"))
if secG is None:
    secG = OxmlElement("w:pgMar"); sec._sectPr.append(secG)
secG.set(qn("w:footer"), str(int(Mm(28).emu / 12700)))  # twips

# ---------- 工具函数 ----------
def set_run(r, ea, size, color=None, bold=False):
    r.font.size = Pt(size)
    r.font.bold = bold
    rPr = r._element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rPr.append(rf)
    rf.set(qn("w:ascii"), ea)
    rf.set(qn("w:hAnsi"), ea)
    rf.set(qn("w:eastAsia"), ea)
    if color:
        c = rPr.find(qn("w:color"))
        if c is None:
            c = OxmlElement("w:color"); rPr.append(c)
        c.set(qn("w:val"), color)

import re
def clean_dates(t):
    """清除日期数字与年/月/日之间的多余空格，以及括号前多余空格。"""
    # 数字后空格+年/月/日 → 数字直接+年/月/日
    t = re.sub(r'(\d)\s+(年|月|日)', r'\1\2', t)
    # "） （" → "）（"（括号间多余空格）
    t = re.sub(r'）\s+（', r'）（', t)
    return t

def add_para(text="", align=WD_ALIGN_PARAGRAPH.LEFT, ea=FONT_FS, size=PT3,
             color=None, left_pt=None, right_pt=None, first_chars=None,
             line_rule=WD_LINE_SPACING.EXACTLY, line_pt=28, space_before=0,
             space_after=0, bold=False):
    text = clean_dates(text)  # 自动清理日期空
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = line_rule
    pf.line_spacing = Pt(line_pt)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if left_pt is not None:
        pf.left_indent = Pt(left_pt)
    if right_pt is not None:
        pf.right_indent = Pt(right_pt)
    # 首行缩进 N 字符（用 python-docx 原生 API，确保 Word 正确渲染）
    if first_chars is not None:
        pf.first_line_indent = Pt(first_chars * size)
    if text:
        r = p.add_run(text)
        set_run(r, ea, size, color, bold)
    return p

def add_red_separator():
    p = add_para("", line_pt=4, space_after=0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")       # 1pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), RED)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def page_break():
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._element.append(br)

def add_page_field(align, color=RED if False else None):
    # 页脚段落：— N —（宋体四号半角）
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(14)
    r1 = p.add_run("— ")
    set_run(r1, FONT_SONG, PT4)
    # PAGE 域
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts"); rf.set(qn("w:ascii"), FONT_SONG); rf.set(qn("w:hAnsi"), FONT_SONG); rf.set(qn("w:eastAsia"), FONT_SONG)
    rPr.append(rf)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(PT4*2)); rPr.append(sz)
    t = OxmlElement("w:t"); t.text = "1"
    r.append(rPr); r.append(t)
    fld.append(r)
    p._p.append(fld)
    r2 = p.add_run(" —")
    set_run(r2, FONT_SONG, PT4)
    return p

# ===================== 正文构建 =====================
# H1 发文机关标志（红头）
add_para(ORG_NAME, align=WD_ALIGN_PARAGRAPH.CENTER,
         ea=FONT_XBS, size=PT1, color=RED, line_rule=WD_LINE_SPACING.AT_LEAST, line_pt=48)
# H3 红色分隔线
add_red_separator()
# B1 标题：分隔线下空二行
add_para("", line_pt=28)
add_para("", line_pt=28)
add_para(DOC_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER,
         ea=FONT_XBS, size=PT2, line_rule=WD_LINE_SPACING.AT_LEAST, line_pt=28)
# 标题下空一行 → 正文
add_para("", line_pt=28)

# 主体正文（逐段）
# 一、工作总体要求（一级 黑体）
add_para("一、工作总体要求", ea=FONT_HEI, size=PT3)
add_para("（一）工作目标", ea=FONT_KAI, size=PT3)  # 二级 楷体
add_para("以 2026 年 6 月 30 日为清查基准日，全面摸清本公司及全资子公司全部国有资产底数，排查闲置低效、权属不清、长期挂账等问题，分类盘活存量资产，补齐管理漏洞，严防国资流失，实现资产规范化长效管理。", first_chars=2)
add_para("（二）工作原则", ea=FONT_KAI, size=PT3)
add_para("全面清查、不留死角；", first_chars=2)
add_para("依规操作、守住底线；", first_chars=2)
add_para("分类盘活、提质增效；", first_chars=2)
add_para("提速推进、边查边改。", first_chars=2)
add_para("（三）清查范围", ea=FONT_KAI, size=PT3)
add_para("厂房、生产设备、存货、车辆等固定资产，土地使用权、软件等无形资产，货币资金、应收账款、预付账款、其他应收款等流动资产，在建工程、账外资产等。", first_chars=2)

# 二、组织架构（一级）
add_para("二、组织架构", ea=FONT_HEI, size=PT3)
add_para("成立资产清查盘活领导小组，总经理组长，总经理助理、财务负责人任副组长； 下设工作专班，财务部牵头，生产、设备、仓储、综合管理协同开展工作，各单位负责人为第一责任人。", first_chars=2)

# 三、工作步骤（一级）
add_para("三、工作步骤", ea=FONT_HEI, size=PT3)
add_para("（一）动员部署（2026年7月10日前完成）", ea=FONT_KAI, size=PT3)
add_para("召开专项工作会，明确分工、标准、时限；", first_chars=2)
add_para("制定简易盘点表、资产台账、问题整改台账、闲置资产盘活台账。", first_chars=2)
add_para("（二）全面清查 （2026 年 7 月底前完成）", ea=FONT_KAI, size=PT3)
add_para("账务核对：以 6 月 30 日为节点清理总账、往来账款，修正错账乱账，做到账账相符；", first_chars=2)
add_para("实地盘点：设备逐台清点、存货全盘、厂房车辆现场核验，核查权属、抵押、闲置、损毁情况；", first_chars=2)
add_para("汇总所有资产数据、问题清单，7 月底前完成资料上报。", first_chars=2)
add_para("（三）分类盘活处置（2026 年 10 月底前完成）", ea=FONT_KAI, size=PT3)
add_para("按 “一资一方案” 分类处置低效闲置资产：", first_chars=2)
add_para("闲置厂房 / 设备：内部调剂、对外出租、技改升级；", first_chars=2)
add_para("积压存货：折价清库、抵债处置；", first_chars=2)
add_para("逾期应收款：法务催收、诉讼清收。", first_chars=2)
add_para("（四）整改建章立制（2026 年 11 月底前完成）", ea=FONT_KAI, size=PT3)
add_para("建立问题整改销号台账，全部问题闭环整改；", first_chars=2)
add_para("完善设备、存货、应收款、闲置资产管理制度；", first_chars=2)
add_para("建立常态化盘点、动态资产监管长效机制。", first_chars=2)

# B5 附件说明（左空二字，名称无标点）
add_para(f"附件：1. {LEADER_GROUP}",
         ea=FONT_FS, size=PT3, left_pt=PT3*2)
# B6 署名（右空二字）/ 日期（右空四字）
add_para("", line_pt=28)  # 附件说明下空一行
add_para(ORG_NAME, align=WD_ALIGN_PARAGRAPH.RIGHT,
         ea=FONT_FS, size=PT3, right_pt=PT3*2)
add_para(DOC_DATE, align=WD_ALIGN_PARAGRAPH.RIGHT,
         ea=FONT_FS, size=PT3, right_pt=PT3*4)

# B8 附件另面
page_break()
# 附件页第一行：附件（黑体顶格）
add_para("附件", ea=FONT_HEI, size=PT3, align=WD_ALIGN_PARAGRAPH.LEFT)
# 空一行 → 附件标题居中于版心第三行
add_para("", line_pt=28)
add_para(LEADER_GROUP,
         align=WD_ALIGN_PARAGRAPH.CENTER, ea=FONT_XBS, size=PT2,
         line_rule=WD_LINE_SPACING.AT_LEAST, line_pt=28)
add_para("", line_pt=28)  # 标题下空一行
add_para("现将工作领导小组人员安排明确如下：", first_chars=2)
add_para("组      长：乔红涛  九天公司总经理", first_chars=2)
add_para("常务副组长：刘  楠  九天公司总经理助理", first_chars=2)
add_para("副  组  长：乔玉娇  九天公司财务总监", first_chars=2)
add_para("成      员：", first_chars=2)
add_para("庞  可   九天公司生产总监", first_chars=2)
add_para("田金梅   九天公司综合管理部主任", first_chars=2)
add_para("唐彩霞   九天公司仓储部人员", first_chars=2)
add_para("郑金玲   九天公司财务部人员", first_chars=2)
add_para("张梦凡   九天公司财务部人员", first_chars=2)
add_para("任  洁   九天公司财务部人员", first_chars=2)

# ===================== 页码（奇偶页不同）PG =====================
def build_footer_p(align):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "right" if align == WD_ALIGN_PARAGRAPH.RIGHT else "left"); pPr.append(jc)
    sp = OxmlElement("w:spacing"); sp.set(qn("w:line"), str(int(Pt(14) / Pt(1) * 240))); sp.set(qn("w:lineRule"), "exact"); pPr.append(sp)
    p.append(pPr)
    def add_run(text):
        r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
        rf = OxmlElement("w:rFonts"); rf.set(qn("w:ascii"), FONT_SONG); rf.set(qn("w:hAnsi"), FONT_SONG); rf.set(qn("w:eastAsia"), FONT_SONG); rPr.append(rf)
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(PT4*2)); rPr.append(sz)
        t = OxmlElement("w:t"); t.text = text; r.append(rPr); r.append(t); p.append(r)
    add_run("— ")
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts"); rf.set(qn("w:ascii"), FONT_SONG); rf.set(qn("w:hAnsi"), FONT_SONG); rf.set(qn("w:eastAsia"), FONT_SONG); rPr.append(rf)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(PT4*2)); rPr.append(sz)
    t = OxmlElement("w:t"); t.text = "1"; r.append(rPr); r.append(t); fld.append(r); p.append(fld)
    add_run(" —")
    return p

def set_footer_content(footer_obj, align):
    ftr = footer_obj._element
    for p in ftr.findall(qn("w:p")):
        ftr.remove(p)
    ftr.append(build_footer_p(align))

# 启用奇偶页不同
sectPr = sec._sectPr
eoh = sectPr.find(qn("w:evenAndOddHeaders"))
if eoh is None:
    eoh = OxmlElement("w:evenAndOddHeaders"); eoh.set(qn("w:val"), "true"); sectPr.append(eoh)
# 奇数页（默认）页脚：居右
set_footer_content(sec.footer, WD_ALIGN_PARAGRAPH.RIGHT)
# 偶数页页脚：居左
even_part, rId = sec.part.add_footer_part()
fr = OxmlElement("w:footerReference")
fr.set(qn("w:type"), "even")
fr.set(qn("w:id"), rId)
sectPr.append(fr)
set_footer_content(even_part, WD_ALIGN_PARAGRAPH.LEFT)

doc.save(OUT)
print("SAVED:", OUT, os.path.exists(OUT))
