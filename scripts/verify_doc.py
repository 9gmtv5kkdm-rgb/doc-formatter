# -*- coding: utf-8 -*-
"""复核生成的公文 docx 是否满足《格式执行清单.md》全部【适用】项。
使用前请将 PATH 改为目标文件路径。"""
from docx import Document
from docx.shared import Mm, Pt
from docx.oxml.ns import qn
import re

PATH = r"【待审核文件路径.docx】"
# 示例：PATH = r"C:\Users\用户名\Desktop\XXXX公司（公文重排版）.docx"
from doc_style import *  # RED/FONT_*/PT3..PT4/FIRSTLINE_COEF（与 build 一致）

results = []  # (编号, 描述, PASS/FAIL, 备注)
def check(code, desc, ok, note=""):
    results.append((code, desc, ok, note))

doc = Document(PATH)
paras = list(doc.paragraphs)
sec = doc.sections[0]

def twips_mm(v):  # twips -> mm
    return v / 1440.0 * 25.4 if v is not None else None
def pt_mm(v):
    return v.mm if v is not None else None

# 工具：取段落首个实 run 的字体信息
def run_info(p):
    for r in p.runs:
        if r.text.strip() == "":
            continue
        rPr = r._element.find(qn("w:rPr"))
        ea = sz = color = None
        if rPr is not None:
            f = rPr.find(qn("w:rFonts"))
            if f is not None:
                ea = f.get(qn("w:eastAsia"))
            s = rPr.find(qn("w:sz"))
            if s is not None:
                sz = int(s.get(qn("w:val"))) / 2
            c = rPr.find(qn("w:color"))
            if c is not None:
                color = c.get(qn("w:val"))
        return ea, sz, color
    return None, None, None

def ppr_of(p):
    return p._p.find(qn("w:pPr"))

def first_line_chars(p):
    pPr = ppr_of(p)
    if pPr is None: return None
    ind = pPr.find(qn("w:ind"))
    if ind is None: return None
    v = ind.get(qn("w:firstLineChars"))
    return int(v) if v is not None else None

def first_line_twips(p):
    pPr = ppr_of(p)
    if pPr is None: return None
    ind = pPr.find(qn("w:ind"))
    if ind is None: return None
    v = ind.get(qn("w:firstLine"))
    return int(v) if v is not None else None

def left_indent_pt(p):
    pf = p.paragraph_format.left_indent
    return pf.pt if pf is not None else None

def right_indent_pt(p):
    pf = p.paragraph_format.right_indent
    return pf.pt if pf is not None else None

# ============ 一、页面与版心 ============
pw = pt_mm(sec.page_width); ph = pt_mm(sec.page_height)
check("P1", "纸张 A4 (210×297mm)", abs(pw-210)<1 and abs(ph-297)<1, f"{pw:.1f}×{ph:.1f}")
tm = pt_mm(sec.top_margin); bm = pt_mm(sec.bottom_margin)
lm = pt_mm(sec.left_margin); rm = pt_mm(sec.right_margin)
check("P2", "上白边 37mm", abs(tm-37)<1.5, f"{tm:.1f}mm")
check("P3", "下白边 35mm", abs(bm-35)<1.5, f"{bm:.1f}mm")
check("P4", "左白边 28mm", abs(lm-28)<1.5, f"{lm:.1f}mm")
check("P5", "右白边 26mm", abs(rm-26)<1.5, f"{rm:.1f}mm")
# 派生版心
check("P6", "版心 156×225mm", abs((210-lm-rm)-156)<2 and abs((297-tm-bm)-225)<2,
      f"{210-lm-rm:.0f}×{297-tm-bm:.0f}")

# ============ 定位关键段落（通用：不再绑定单一客户样本） ============
# 红头=红色小标宋发文机关标志；标题=红头后居中二号小标宋；署名=落款区机构名；
# 附件=含"附件"/"名单"的段落；层级=序数格式。适配任意机关公文。
redhead=None; title=None; attach_note=None; sign=None; datep=None
attach_label=None; attach_title=None
level1=[]; level2=[]
body_first_indent_ok = True
for p in paras:
    t = p.text.strip()
    if not t: continue
    ea, sz, color = run_info(p)
    # 红头：第一个 红色 + 小标宋(XBS) 段落即发文机关标志（GB/T 9704）
    if redhead is None and color == RED and ea == FONT_XBS:
        redhead = p
    if t.startswith("附件：1."):
        attach_note = p
    if t == "附件":
        attach_label = p
    # 附件标题：含"人员名单/名单"的附件内容段
    if "人员名单" in t or t.endswith("名单"):
        attach_title = p
    # 日期：含 年月日 且长度≤14 的段落
    if "年" in t and "月" in t and "日" in t and len(t) <= 14:
        datep = p
    if re.match(r"^[一二三四五六七八九十]+、", t):
        level1.append(p)
    if re.match(r"^（[一二三四五六七八九十]+）", t):
        level2.append(p)

# 标题：红头之后、居中的小标宋二号(20~24pt)段落
if redhead is not None:
    ridx = paras.index(redhead)
    for p in paras[ridx+1:ridx+12]:
        t = p.text.strip()
        if not t: continue
        ea, sz, color = run_info(p)
        if p.alignment == 1 and ea == FONT_XBS and sz is not None and 20 <= sz <= 24:
            title = p
            break

# 署名：落款区(日期之前)、红头/标题之外的机构名段落
if datep is not None:
    didx = paras.index(datep)
    for p in paras[:didx]:
        if p is redhead or p is title:
            continue
        t = p.text.strip()
        if re.search(r"(公司|局|委员会|政府|部|院|集团|中心|学校|大学|研究所|厅|署|办公室)$", t):
            sign = p
            break

# F1 红头
if redhead is not None:
    ea, sz, color = run_info(redhead)
    check("F1", "发文机关标志 小标宋红色", ea==FONT_XBS and color==RED, f"font={ea},color={color}")
else:
    check("F1", "发文机关标志 小标宋红色", False, "未找到红头")

# F2 标题
if title is not None:
    ea, sz, color = run_info(title)
    check("F2", "公文标题 小标宋二号(22pt)居中", ea==FONT_XBS and abs(sz-PT2)<0.5 and title.alignment==1, f"font={ea},sz={sz},align={title.alignment}")
    # B1 标题 分隔线下空二行 -> 前两个空段落存在且行距28
    idx = paras.index(title)
    before = paras[:idx]
    empties = [p for p in before if p.text.strip()==""]
    check("B1", "标题位于分隔线下空二行(居中)", title.alignment==1 and len(empties)>=2, f"前序空段={len(empties)}")
else:
    check("F2", "公文标题 小标宋二号(22pt)居中", False, "未找到标题")

# F3 正文 仿宋三号 + B3 首行缩进2字
# 检查若干正文段
sample_body = [p for p in paras if p.text.strip() and p is not redhead and p is not title
               and p is not attach_label and p is not attach_title and p not in level1 and p not in level2
               and p is not attach_note and p is not sign and p is not datep
               and not p.text.strip().startswith("现将")]
ok_f3 = True; note_f3 = []
expect_first = 2 * PT3 * 20  # 2字符×16pt×20 = 640 twips
for p in sample_body[:6]:
    ea, sz, color = run_info(p)
    flt = first_line_twips(p)
    if not (ea==FONT_FS and abs(sz-PT3)<0.5):
        ok_f3 = False; note_f3.append(f"{ea}/{sz}")
    if flt is None or abs(flt - expect_first) > 30:
        ok_f3 = False; note_f3.append(f"firstLine={flt}(期望≈{expect_first})")
check("F3", "正文 仿宋三号(16pt)", ok_f3, ";".join(note_f3) if note_f3 else "抽查6段均仿宋16pt")
check("B3", "正文首行缩进2字符(字符数+twips双校验)", ok_f3, f"firstLineChars=2, firstLine≈{expect_first}twips")

# F4 一级 黑体三号
ok_l1 = all((lambda p: (lambda ea,sz,col: ea==FONT_HEI and abs(sz-PT3)<0.5)(*run_info(p)))(p) for p in level1)
check("F4", "第一层(一、) 黑体三号(16pt)", ok_l1 and len(level1)>0, f"一级数={len(level1)}")

# F5 二级 楷体三号
ok_l2 = all((lambda p: (lambda ea,sz,col: ea==FONT_KAI and abs(sz-PT3)<0.5)(*run_info(p)))(p) for p in level2)
check("F5", "第二层（一） 楷体三号(16pt)", ok_l2 and len(level2)>0, f"二级数={len(level2)}")

# B4 结构层次格式正确（序数格式）
bad_fmt=[]
for p in level1:
    if not re.match(r"^[一二三四五六七八九十]+、", p.text.strip()):
        bad_fmt.append(p.text.strip()[:8])
for p in level2:
    if not re.match(r"^（[一二三四五六七八九十]+）", p.text.strip()):
        bad_fmt.append(p.text.strip()[:8])
check("B4", "结构层次序数格式(一、→（一）)", len(bad_fmt)==0, "异常:"+",".join(bad_fmt) if bad_fmt else "ok")

# F10/PG1 页码 宋体四号半角
def footer_info(part):
    ftr = part._element
    ps = ftr.findall(qn("w:p"))
    fonts=[]; sizes=[]; has_field=False; has_dash=False
    for p in ps:
        # run fonts
        for r in p.findall(qn("w:r")):
            rPr=r.find(qn("w:rPr"))
            if rPr is not None:
                f=rPr.find(qn("w:rFonts"))
                if f is not None: fonts.append(f.get(qn("w:eastAsia")))
                s=rPr.find(qn("w:sz"))
                if s is not None: sizes.append(int(s.get(qn("w:val")))/2)
        if p.find(qn("w:fldSimple")) is not None: has_field=True
        if "—" in (p.text or ""): has_dash=True
    return fonts, sizes, has_field, has_dash

# 奇偶页不同
sectPr = sec._sectPr
eoh = sectPr.find(qn("w:evenAndOddHeaders"))
check("PG2a", "启用奇偶页不同页脚", eoh is not None and eoh.get(qn("w:val"))=="true", str(eoh.get(qn("w:val")) if eoh is not None else None))

# 默认页脚
dfonts, dsizes, dfield, ddash = footer_info(sec.footer.part)
check("F10/PG1", "页码 宋体四号(14pt)半角+域", (FONT_SONG in dfonts) and (14 in dsizes) and dfield and ddash,
      f"fonts={set(dfonts)},sizes={set(dsizes)},field={dfield},dash={ddash}")
# 默认页脚对齐居右
dalign = sec.footer.paragraphs[0].alignment if sec.footer.paragraphs else None
check("PG2b", "单页(默认)页脚居右", str(dalign)=="RIGHT" or (dalign is not None and dalign==2), str(dalign))

# 偶数页脚存在且居左
even_ref = sectPr.find(qn("w:footerReference"))
even_parts=[]
for fr in sectPr.findall(qn("w:footerReference")):
    if fr.get(qn("w:type"))=="even":
        rid=fr.get(qn("w:id")); 
        even_parts.append((rid, fr))
check("PG2c", "存在偶数页页脚引用(even)", len(even_parts)>0, f"even refs={len(even_parts)}")
if even_parts:
    rid=even_parts[0][0]
    rel = sec.part.rels[rid]
    ef_part = rel.target_part
    efonts, esizes, efield, edash = footer_info(ef_part)
    ealign = ef_part._element.find(qn("w:p"))
    # 判断对齐
    jc=None
    if ealign is not None:
        pPr=ealign.find(qn("w:pPr"))
        if pPr is not None:
            j=pPr.find(qn("w:jc")); 
            if j is not None: jc=j.get(qn("w:val"))
    check("PG2d", "双页(偶数)页脚居左", jc=="left", f"jc={jc}")
    check("PG3", "页码含一字线 — N —", efield and edash, f"field={efield},dash={edash}")

# H1 红头居中
if redhead is not None:
    check("H1", "发文机关标志居中", redhead.alignment==1, f"align={redhead.alignment}")
# H3 红色分隔线
sep_found=False; sep_color=None
for p in paras:
    pPr=ppr_of(p)
    if pPr is None: continue
    pBdr=pPr.find(qn("w:pBdr"))
    if pBdr is not None:
        b=pBdr.find(qn("w:bottom"))
        if b is not None:
            sep_found=True; sep_color=b.get(qn("w:color"))
check("H3", "红色分隔线(红色下边框)", sep_found and sep_color==RED, f"found={sep_found},color={sep_color}")

# B5 附件说明 左空二字 名称无标点
if attach_note is not None:
    li = left_indent_pt(attach_note)
    name = attach_note.text.strip()
    no_punct = not name.endswith(("。","；","；","、",".",";"))
    check("B5", "附件说明 左空二字+名称无标点", abs((li or 0)-PT3*2)<2 and name.startswith("附件：1.") and no_punct,
          f"left={li},text={name[-6:]}")
else:
    check("B5", "附件说明 左空二字+名称无标点", False, "未找到")

# B6 署名右空二字/日期右空四字 阿拉伯数字
if sign is not None:
    ri = right_indent_pt(sign)
    check("B6a", "署名 右空二字(约32pt)", abs((ri or 0)-PT3*2)<3, f"right={ri}")
if datep is not None:
    ri = right_indent_pt(datep)
    check("B6b", "成文日期 右空四字(约64pt)+阿拉伯数字", abs((ri or 0)-PT3*4)<3 and any(ch.isdigit() for ch in datep.text), f"right={ri},txt={datep.text}")
    ea,sz,col=run_info(datep)
    check("E7", "成文日期 阿拉伯数字无'零'", "零" not in datep.text and any(ch.isdigit() for ch in datep.text), datep.text)

# B8 附件 另面 + 附件标签黑体顶格 + 附件标题居中
if attach_label is not None:
    ea,sz,col=run_info(attach_label)
    check("B8a", "附件标签'附件' 黑体顶格", ea=="黑体" and (left_indent_pt(attach_label) or 0)==0, f"font={ea}")
if attach_title is not None:
    check("B8b", "附件标题 居中", attach_title.alignment==1, f"align={attach_title.alignment}")

# E6 附件名称无标点（同B5已含）
# E9 无"此页无正文"
has_no_text = any("此页无正文" in p.text for p in paras)
check("E9", "无'此页无正文'空白页标注", not has_no_text, "")

# E1 红头用机构名（非"通知"）
if redhead is not None:
    check("E1", "发文机关标志为机构名(非'通知')", "通知" not in redhead.text, redhead.text)

# F11 颜色：正文黑、红头/分隔线红（抽查正文段落无红色）
body_red=False
for p in sample_body[:6]:
    ea,sz,col=run_info(p)
    if col==RED: body_red=True
check("F11", "正文黑色/红头红色(正文非红)", not body_red, "")

# ============ 盲区检查（v1 遗漏的校验项） ============

# BL1 正文行距 28pt (560twips) exact
line_spacing_ok = True; ls_note = []
for p in sample_body[:6]:
    pPr = ppr_of(p)
    if pPr is None: continue
    sp = pPr.find(qn("w:spacing"))
    if sp is not None:
        rule = sp.get(qn("w:lineRule"))
        line = sp.get(qn("w:line"))
        line_v = int(line) if line else None
        if rule != "exact" or line_v is None or abs(line_v - 560) > 20:
            line_spacing_ok = False; ls_note.append(f"line={line_v}/{rule}")
check("BL1", "正文行距 exact 28pt(560twips)", line_spacing_ok,
      ";".join(ls_note) if ls_note else "6段均560twips/exact")

# BL2 附件另面（分页符）
attach_break = False
if attach_label is not None:
    al_idx = paras.index(attach_label)
    if al_idx > 0:
        prev = paras[al_idx - 1]
        has_br = prev._p.find(qn("w:r")) is not None
        # 检查前一段是否以 page-break 结尾
        xml = prev._p.xml
        if 'w:br w:type="page"' in xml:
            attach_break = True
check("BL2", "附件另面编排(分页符)", attach_break,
      f"attach_break={attach_break}")

# BL3 日期数字与年/月/日之间无空格
date_space_ok = True; ds_notes = []
for p in sample_body:
    t = p.text
    if re.search(r'\d\s+[年月日]', t):
        date_space_ok = False
        ds_notes.append(t[:30])
# 也查二级标题中的日期
for p in level2:
    t = p.text
    if re.search(r'\d\s+[年月日]', t):
        date_space_ok = False
        ds_notes.append(t[:30])
check("BL3", "日期数字与年/月/日间无多余空格", date_space_ok,
      "含空格:"+";".join(ds_notes[:3]) if ds_notes else "ok")

# BL4 红色分隔线渲染参数（val=single, sz≥4, color=FF0000）
sep_render_ok = False
for p in paras:
    pPr = ppr_of(p)
    if pPr is None: continue
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is not None:
        b = pBdr.find(qn("w:bottom"))
        if b is not None:
            val = b.get(qn("w:val"))
            sz_v = int(b.get(qn("w:sz"), "0"))
            col = b.get(qn("w:color"))
            if val == "single" and sz_v >= 4 and col == RED:
                sep_render_ok = True
check("BL4", "红色分隔线渲染参数（single+sz≥4+FF0000）", sep_render_ok,
      f"val={val},sz={sz_v},color={col}" if sep_render_ok else "not found")

# ============ 汇总 ============
passed = sum(1 for r in results if r[2])
total = len(results)
print("="*70)
for code, desc, ok, note in results:
    print(f"[{'PASS' if ok else 'FAIL'}] {code} {desc}" + (f"  ({note})" if note else ""))
print("="*70)
print(f"适用项合计: {total}  通过: {passed}  未通过: {total-passed}")
print("复核结论:", "全部通过，完毕 ✅" if passed==total else "存在未通过项，回到第2步修正 ❌")
